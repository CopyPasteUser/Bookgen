#!/usr/bin/env python3
from g4f.client import Client
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tqdm import tqdm
import time
import threading
from queue import Queue
import argparse

client = Client()

def askAI(message, max_retries=3):
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model="gpt-4.1",
                messages=[{"role": "user", "content": message}],
                web_search=False
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"Attempt {attempt + 1} failed: {e}")
            if attempt == max_retries - 1:
                print("Could not reach the AI service. Please try again later.")
                return None

def get_book_structure(topic, competence_level, language, max_chapters=12):
    chapter_keyword = "chapters" if language.lower() == "english" else "kapitel"
    chapter_title_key = "chapter_title" if language.lower() == "english" else "chapter_title"
    
    prompt = f"""Create a book structure in JSON format about '{topic}' for {competence_level} level readers.
The book should have between 3 and {max_chapters} chapters.
Each chapter should have 3-5 subchapters.
Use this exact JSON format:
{{
    "book_title": "Book Title",
    "{chapter_keyword}": [
        {{
            "{chapter_title_key}": "Chapter 1 Title",
            "subchapters": ["Subchapter 1.1", "Subchapter 1.2", "Subchapter 1.3"]
        }}
    ]
}}
Write in {language} language. Respond ONLY with valid JSON."""
    
    response = askAI(prompt)
    try:
        return json.loads(response)
    except (json.JSONDecodeError, TypeError):
        print("Error parsing JSON response")
        return None

def generate_chapter_content(topic, chapter_title, subchapters, competence_level, language, result_queue, index):
    content = f"# {chapter_title}\n\n"
    
    for subchapter in subchapters:
        prompt = f"""Write a detailed section about '{subchapter}' for a book about '{topic}'.
The book is for {competence_level} level readers.
Write in {language} language with appropriate complexity.
Add examples and practical applications where relevant.
Section length: 300-500 words."""
        
        subchapter_content = askAI(prompt)
        if subchapter_content is None:
            result_queue.put((index, None))
            return
            
        content += f"## {subchapter}\n\n{subchapter_content}\n\n"
    
    result_queue.put((index, content))

def create_manual_toc(doc, chapters, language):
    toc_title = "Table of Contents" if language.lower() == "english" else "Inhaltsverzeichnis"
    description = "Here you'll find an overview of all chapters:" if language.lower() == "english" else "Hier finden Sie eine Übersicht aller Kapitel:"
    chapter_word = "Chapter" if language.lower() == "english" else "Kapitel"
    
    doc.add_heading(toc_title, level=0)
    doc.add_paragraph(description, style='Intense Quote')
    
    for i, chapter in enumerate(chapters, 1):
        # Chapter line
        p = doc.add_paragraph()
        p.add_run(f"{chapter_word} {i}: ").bold = True
        p.add_run(chapter['chapter_title']).italic = True
        
        # Subchapters
        for j, subchapter in enumerate(chapter['subchapters'], 1):
            doc.add_paragraph(f"    {i}.{j} {subchapter}", style='List Bullet')
    
    doc.add_page_break()

def markdown_to_docx(md_text, output_path, structure, language):
    doc = Document()
    
    # Document styles
    styles = doc.styles
    font = styles['Normal'].font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title page
    lines = md_text.split('\n')
    if lines:
        title = lines[0].lstrip('#').strip()
        if title:
            # Title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Metadata
            doc.add_paragraph("\n" * 3)  # Spacing
            meta = doc.add_paragraph()
            generated_text = "Generated on: " if language.lower() == "english" else "Generiert am: "
            meta.add_run(generated_text + datetime.now().strftime("%d.%m.%Y")).italic = True
            meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            doc.add_page_break()
    
    # Table of contents
    create_manual_toc(doc, structure['chapters'], language)
    
    # Main content
    current_chapter = 0
    chapter_word = "Chapter" if language.lower() == "english" else "Kapitel"
    
    for line in md_text.split('\n'):
        if line.startswith('# '):
            current_chapter += 1
            heading = doc.add_heading(f"{chapter_word} {current_chapter}: {line[2:].strip()}", level=1)
            heading.paragraph_format.space_before = Pt(18)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip() and not line.startswith('*'):
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Inches(0.25)
            p.paragraph_format.line_spacing = 1.15
    
    # Headers and footers
    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = structure['book_title']
        paragraph.style = doc.styles['Header']
        
        footer = section.footer
        paragraph = footer.paragraphs[0]
        page_text = "Page \\p of \\P" if language.lower() == "english" else "Seite \\p von \\P"
        paragraph.text = page_text
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.save(output_path)

def create_book(topic, competence_level, language, output_dir="books"):
    os.makedirs(output_dir, exist_ok=True)
    
    print("\n📖 Generating book contents..." if language.lower() == "english" else "\n📖 Bucheinhalte werden generiert...")
    
    # Get book structure
    with tqdm(total=100, desc="Creating structure" if language.lower() == "english" else "Struktur erstellen", 
              bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt}") as pbar:
        structure = get_book_structure(topic, competence_level, language)
        pbar.update(30)
        if not structure:
            print("Error: Could not create book structure" if language.lower() == "english" else "Fehler: Buchstruktur konnte nicht erstellt werden")
            return
    
    # Generate filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    sanitized_topic = "".join(c if c.isalnum() else "_" for c in topic)
    filename_base = f"{output_dir}/{sanitized_topic}_{timestamp}"
    
    # Create markdown content
    md_content = f"# {structure['book_title']}\n\n"
    md_content += f"*Topic: {topic}*\n" if language.lower() == "english" else f"*Thema: {topic}*\n"
    md_content += f"*Level: {competence_level}*\n"
    md_content += f"*Language: {language}*\n\n" if language.lower() == "english" else f"*Sprache: {language}*\n\n"
    
    # Threaded chapter generation
    result_queue = Queue()
    threads = []
    chapter_contents = [None] * len(structure['chapters'])
    
    with tqdm(total=len(structure['chapters']), 
              desc="Generating chapters" if language.lower() == "english" else "Kapitel generieren",
              bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt}") as pbar:
        
        for i, chapter in enumerate(structure['chapters']):
            thread = threading.Thread(
                target=generate_chapter_content,
                args=(topic, chapter["chapter_title"], chapter["subchapters"], 
                      competence_level, language, result_queue, i)
            )
            threads.append(thread)
            thread.start()
            time.sleep(0.5)  # Rate limiting
        
        # Wait for threads to complete
        for _ in range(len(structure['chapters'])):
            idx, content = result_queue.get()
            if content is None:
                print("\nError generating content for one chapter. Stopping.")
                for t in threads:
                    t.join()
                return
            chapter_contents[idx] = content
            pbar.update(1)
        
        for thread in threads:
            thread.join()
    
    # Combine all chapter contents
    md_content += "\n\n".join(chapter_contents)
    
    # Create DOCX file
    with tqdm(total=1, desc="Creating Word document" if language.lower() == "english" else "Word-Dokument erstellen",
              bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt}") as pbar:
        docx_filename = f"{filename_base}.docx"
        markdown_to_docx(md_content, docx_filename, structure, language)
        pbar.update(1)
    
    success_msg = "\n✅ Book successfully created!" if language.lower() == "english" else "\n✅ Buch erfolgreich erstellt!"
    print(success_msg)
    print(f"📄 Word document: {os.path.abspath(docx_filename)}" if language.lower() == "english" else f"📄 Word-Dokument: {os.path.abspath(docx_filename)}")

def main():
    parser = argparse.ArgumentParser(description='AI Book Generator')
    parser.add_argument('--topic', help='Book topic')
    parser.add_argument('--level', help='Competence level (e.g., Beginner, Advanced)')
    parser.add_argument('--lang', help='Language (e.g., english, deutsch)')
    
    args = parser.parse_args()
    
    if not all([args.topic, args.level, args.lang]):
        print("📚 AI Book Generator v1.0" if args.lang and args.lang.lower() == "english" else "📚 AI-Buchgenerator v1.0")
        print("-----------------------\n")
        
        if not args.topic:
            args.topic = input("Book topic: " if args.lang and args.lang.lower() == "english" else "Thema des Buches: ")
        if not args.level:
            args.level = input("Difficulty level: " if args.lang and args.lang.lower() == "english" else "Schwierigkeitsgrad: ")
        if not args.lang:
            args.lang = input("Language (english/deutsch): ").lower()
    
    create_book(args.topic, args.level, args.lang)

if __name__ == "__main__":
    main()
